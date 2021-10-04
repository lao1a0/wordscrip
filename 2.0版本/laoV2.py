'''
@Time : 2021-10-02 20:05
@Author : laolao
@FileName: laoV2.py
查词来源：https://www.jisuapi.com/api/cidian/,用的是一个汉语api,免费的是100次/天,用的时候把API改成自己的
'''
import json

import dns
from docx import Document
from docx.shared import Pt, Cm
from docx.shared import Inches
from docx.oxml.ns import qn
from docx import Document
from docx.enum.section import WD_ORIENT
import xlrd
import xlwt
from lxml import etree
import requests
import urllib
import re
API ='5ad8766465d53597'
def formats(str):
    pattern = re.compile(r'<[^>]+>',re.S)
    if str !=None:
        return pattern.sub('',  str)
    else:
        return str
def get_explain_for_each_word(excel_file_name):
    '''
    输入excel文件的名字，文件用orc识别出来的，每个词语一列
    :return:词语，词语的意思，来源网站
    '''
    readbook = xlrd.open_workbook(excel_file_name)
    sheet = readbook.sheet_by_index(0)
    nrows = sheet.nrows # 行
    headers = {
        'User-Agent': "Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.103 Safari/537.36",
    }
    word_list=[]
    for i in range(nrows):
        cel = sheet.row_values(i)[0].split('.') # 获取i行表格值
        if len(cel) == 2:
            line = cel[1].strip()
            zh = "".join(re.compile('[^\u4e00-\u9fa5]').split(line)).strip() # 中文的编码范围是：\u4e00到\u9fa5
            _url = "https://api.jisuapi.com/cidian/word?appkey={}&word={}".format(API,urllib.parse.quote(zh))
            req = requests.get(url=_url ,headers=headers).json()
            if req['msg']=="ok":
                    name = req['result']['name']
                    pinyi = req['result']['pinyin']
                    content =formats(req['result']['content'])
                    basiccontent = formats(req['result']['basiccontent'])
                    example = formats(req['result']['example'])
                    comefrom = formats(req['result']['comefrom'])
                    english = req['result']['english']
                    jin =req['result']['jin']
                    fan =req['result']['fan']
                    word_list.append([name,pinyi,basiccontent,content,example,comefrom,english,jin,fan])
            else:
                word_list.append([zh,"","","","","","","",""])
    return word_list

def return_to_excel(day,dir):
    workbook = xlwt.Workbook(encoding = 'utf-8')
    worksheet = workbook.add_sheet("sheet1")
    # 参数对应 行, 列, 值
    worksheet.write(0,0, label ='词语')
    worksheet.write(0,1, label ='拼音')
    worksheet.write(0,2, label ='基本意思')
    worksheet.write(0,3, label ='内容')
    worksheet.write(0,4, label ='例句')
    worksheet.write(0,5, label ='来源')
    worksheet.write(0,6, label ='英文')
    worksheet.write(0,7, label ='近义词')
    worksheet.write(0,8, label ='反义词')
    i=1
    for d in dir:
        name,pinyi,basiccontent,content,example,comefrom,english,jin,fan = d
        # print(d)
        worksheet.write(i,0, label =name)
        worksheet.write(i,1, label =pinyi)
        worksheet.write(i,2, label =basiccontent)
        worksheet.write(i,3, label =content)
        worksheet.write(i,4, label =example)
        worksheet.write(i,5, label =comefrom)
        worksheet.write(i,6, label =english)
        worksheet.write(i,7, label =jin)
        worksheet.write(i,8, label =fan)
        i+=1
    workbook.save('第{}天.xls'.format(day))
def return_to_docx(day,dir):
        '''
        结果输入到word里面
        :param day:
        :param dir:
        :return:
        '''
        #打开文档，设置格式
        document = Document()
        section = document.sections[0]
        # 横向
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        # 窄
        section.top_margin=Cm(1.27)
        section.bottom_margin=Cm(1.27)
        section.left_margin=Cm(1.27)
        section.right_margin=Cm(1.27)
        #加入不同等级的标题
        document.add_heading(u'第{}天'.format(day),1) # 一级标题
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.size = Pt(14)
        for d in dir:
            name,pinyi,basiccontent,content,example,comefrom,english,jin,fan = d
            document.add_heading(u'{}（{}）{}'.format(name,pinyi,english).format(day),2) # 二级标题
            document.add_paragraph().add_run(u'[基本解释]')
            document.add_paragraph().add_run(u'{}'.format(basiccontent))
            document.add_paragraph().add_run(u'[内容]')
            document.add_paragraph().add_run(u'{}'.format(content))
            document.add_paragraph().add_run(u'[例句]')
            document.add_paragraph().add_run(u'{}'.format(example))
            document.add_paragraph().add_run(u'[来源]')
            if comefrom !=None:
                document.add_paragraph().add_run(u'{}'.format(comefrom))
            else:
                document.add_paragraph().add_run(u'{}'.format(''))
            document.add_paragraph().add_run(u'[近义词]')
            document.add_paragraph().add_run(u'{}'.format(jin))
            document.add_paragraph().add_run(u'[反义词]')
            document.add_paragraph().add_run(u'{}'.format(fan))

        #增加分页
        document.add_page_break()
        #保存文件
        document.save(u'第{}天.docx'.format(day))

if __name__ == '__main__':
    dir = get_explain_for_each_word('1.xlsx')
    return_to_docx(100,dir)
    print(">>Finish")